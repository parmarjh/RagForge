"""RAGForge — DOCX document loader."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from ragforge.core.types import Document, SourceType
from ragforge.ingest.base import BaseLoader


class DocxLoader(BaseLoader):
    """Load text content from DOCX files using python-docx."""

    def load(self, source: str, **kwargs: Any) -> list[Document]:
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX loading. "
                "Install it with: pip install python-docx"
            )

        path = Path(source)
        metadata = kwargs.get("metadata", {})

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        full_text = "\n\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n" + "\n".join(table_texts)

        if not full_text.strip():
            return []

        return [
            Document(
                content=full_text,
                source=str(path.resolve()),
                source_type=SourceType.FILE,
                metadata={
                    "filename": path.name,
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                    **metadata,
                },
            )
        ]
